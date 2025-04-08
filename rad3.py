import streamlit as st
import docx
import re
import io
import base64
from typing import Dict, List, Tuple
import pandas as pd

def find_term_values(doc: docx.Document, terms: List[str]) -> Dict[str, List[Tuple[int, str]]]:
    """Find terms and their corresponding values in the document.
    
    Returns: Dictionary with term as key and list of (paragraph_index, value) tuples as value
    """
    results = {term: [] for term in terms}
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text
        for term in terms:
            if term in text:
                # Find a value that appears after the term
                # This is a simplified approach - assuming value is numeric and appears after the term
                value_match = re.search(rf'{term}\s*[:=]?\s*(\d+[\.\d]*)', text)
                if value_match:
                    value = value_match.group(1)
                    results[term].append((i, value))
    
    return results

def replace_and_save_doc(doc: docx.Document, term_values: Dict[str, List[Tuple[int, str]]], 
                         replacements: Dict[str, List[str]]) -> List[bytes]:
    """Replace values in the document and save multiple versions.
    
    Returns: List of document bytes for each replacement set
    """
    # Determine max number of documents to create
    max_replacements = max([len(values) for values in replacements.values()], default=0)
    if max_replacements == 0:
        return []
    
    output_docs = []
    
    for i in range(max_replacements):
        # Instead of creating a new document, we'll use the docx library's deep copy functionality
        # by saving and reloading the document for each variation
        temp_doc_bytes = io.BytesIO()
        doc.save(temp_doc_bytes)
        temp_doc_bytes.seek(0)
        new_doc = docx.Document(temp_doc_bytes)
        
        # Apply replacements to the copied document
        for term, locations in term_values.items():
            if term in replacements and i < len(replacements[term]):
                new_value = replacements[term][i]
                
                # Replace in specific paragraphs where the term was found
                for loc_para_idx, old_value in locations:
                    para = new_doc.paragraphs[loc_para_idx]
                    text = para.text
                    
                    # Replace the value
                    new_text = re.sub(rf'{term}\s*[:=]?\s*{re.escape(old_value)}', f'{term}: {new_value}', text)
                    
                    # If text changed, update paragraph text while preserving formatting
                    if new_text != text:
                        # Save original runs and their formatting
                        original_runs = []
                        for run in para.runs:
                            original_runs.append({
                                'text': run.text,
                                'bold': run.bold,
                                'italic': run.italic,
                                'underline': run.underline,
                                'font_size': run.font.size,
                                'font_color': run.font.color.rgb,
                                'style': run.style
                            })
                        
                        # Clear the paragraph and set new text
                        para.clear()
                        para.add_run(new_text)
                        
                        # If there's only one run in the original and new text, copy formatting
                        if len(original_runs) == 1:
                            for attr, value in original_runs[0].items():
                                if attr != 'text' and value is not None:
                                    if attr == 'font_size':
                                        para.runs[0].font.size = value
                                    elif attr == 'font_color':
                                        para.runs[0].font.color.rgb = value
                                    elif attr == 'style':
                                        para.runs[0].style = value
                                    else:
                                        setattr(para.runs[0], attr, value)
        
        # Save the modified document
        docx_bytes = io.BytesIO()
        new_doc.save(docx_bytes)
        docx_bytes.seek(0)
        output_docs.append(docx_bytes.getvalue())
    
    return output_docs

def create_download_link(file_bytes, filename):
    """Create a download link for a file."""
    b64 = base64.b64encode(file_bytes).decode()
    href = f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="{filename}">Download {filename}</a>'
    return href

def main():
    st.title("Word Document Value Replacer")
    st.write("""
    This app allows you to replace specific values in a Word document with values from a list.
    Upload your Word document, specify the terms to find, and provide replacement values.
    """)
    
    uploaded_file = st.file_uploader("Upload Word Document", type="docx")
    
    if uploaded_file:
        # Load the document
        doc = docx.Document(uploaded_file)
        
        # User inputs terms
        st.subheader("Enter Terms to Replace")
        terms_input = st.text_area("Enter terms to search for (one per line, e.g., LHS, RHS, etc.)")
        
        selected_terms = []
        if terms_input:
            selected_terms = [term.strip() for term in terms_input.strip().split('\n') if term.strip()]
        
        if not selected_terms:
            st.info("Please enter at least one term to search for in the document.")
        
        if selected_terms:
            # Find terms and their values in the document
            term_values = find_term_values(doc, selected_terms)
            
            # Display found terms and values
            st.subheader("Found Values")
            terms_found = False
            for term, locations in term_values.items():
                if locations:
                    terms_found = True
                    st.write(f"**{term}**:")
                    for loc, value in locations:
                        st.write(f"  - Value: {value} (paragraph {loc+1})")
                else:
                    st.write(f"**{term}**: Not found or no value detected")
            
            if not terms_found:
                st.warning("No values were found for the specified terms. Check if the terms exist in the document and have numeric values associated with them.")
            
            # Input for replacement values
            st.subheader("Enter Replacement Values")
            
            # Choose input method
            input_method = st.radio(
                "Choose input method for replacement values:",
                ["Manual Entry", "Upload CSV", "Paste CSV Data"]
            )
            
            replacements = {}
            
            if input_method == "Manual Entry":
                for term, locations in term_values.items():
                    if locations:
                        st.write(f"**{term}** Replacements:")
                        replacement_values = st.text_area(
                            f"Enter replacement values for {term} (one per line)",
                            key=f"replace_{term}"
                        )
                        if replacement_values:
                            replacements[term] = [v.strip() for v in replacement_values.strip().split("\n") if v.strip()]
                
            elif input_method == "Upload CSV":
                csv_file = st.file_uploader("Upload CSV file with replacement values", type="csv")
                if csv_file:
                    df = pd.read_csv(csv_file)
                    st.write("Preview of uploaded data:")
                    st.dataframe(df.head())
                    
                    # Map columns to terms
                    for term in term_values.keys():
                        if term in df.columns:
                            replacements[term] = df[term].dropna().tolist()
                        else:
                            st.warning(f"Column '{term}' not found in CSV")
            
            elif input_method == "Paste CSV Data":
                csv_data = st.text_area("Paste CSV data (with header row)")
                if csv_data:
                    try:
                        data_io = io.StringIO(csv_data)
                        df = pd.read_csv(data_io)
                        st.write("Preview of data:")
                        st.dataframe(df.head())
                        
                        # Map columns to terms
                        for term in term_values.keys():
                            if term in df.columns:
                                replacements[term] = df[term].dropna().tolist()
                            else:
                                st.warning(f"Column '{term}' not found in CSV data")
                    except Exception as e:
                        st.error(f"Error parsing CSV data: {e}")
            
            # Process and generate documents
            if replacements and st.button("Generate Documents"):
                with st.spinner("Generating documents..."):
                    doc_bytes_list = replace_and_save_doc(doc, term_values, replacements)
                    
                    if doc_bytes_list:
                        st.success(f"{len(doc_bytes_list)} documents generated!")
                        
                        # Create download links
                        st.subheader("Download Documents")
                        for i, doc_bytes in enumerate(doc_bytes_list):
                            st.write(f"Document {i+1}:")
                            # Allow custom filename
                            filename = st.text_input(f"Enter filename for document {i+1}", value=f"document_{i+1}.docx", key=f"filename_{i}")
                            st.markdown(create_download_link(doc_bytes, filename), unsafe_allow_html=True)
                            st.write("---")
                    else:
                        st.warning("No documents were generated. Please check your replacement values.")

if __name__ == "__main__":
    main()